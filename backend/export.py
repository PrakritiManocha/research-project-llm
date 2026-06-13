"""
Report export module for the Deep Researcher application.

This module handles the export of research reports in various formats:
1. PDF
2. DOCX
3. HTML
"""

import os
from typing import Dict, Any, List
import logging
from pathlib import Path
from docx import Document as DocxDocument
import markdown
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
import shutil
from datetime import datetime
import re

from backend.config import (
    EXPORTS_DIR,
    PDF_OPTIONS
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ReportExporter:
    """Handles the export of research reports in various formats."""
    
    def __init__(self):
        """Initialize the report exporter."""
        self.exports_dir = Path(EXPORTS_DIR)
        self.exports_dir.mkdir(parents=True, exist_ok=True)
        
        # Define styles
        self.styles = getSampleStyleSheet()
        self.styles.add(ParagraphStyle(
            name='CustomCitation',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceBefore=6,
            spaceAfter=6
        ))
        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceBefore=24,
            spaceAfter=12
        ))
        self.styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=18,
            spaceAfter=8
        ))
        self.styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceBefore=6,
            spaceAfter=6
        ))
    
    def _process_markdown(self, text: str) -> List[str]:
        """Process markdown text into paragraphs."""
        # Split by double newlines to get paragraphs
        paragraphs = text.split('\n\n')
        processed = []
        
        for para in paragraphs:
            # Remove markdown formatting
            para = re.sub(r'\*\*(.*?)\*\*', r'\1', para)  # Remove bold
            para = re.sub(r'#+\s*', '', para)  # Remove headings
            para = re.sub(r'-\s*', '• ', para)  # Convert dashes to bullets
            processed.append(para.strip())
        
        return processed
    
    def export_pdf(self, report: str, citations: List[Dict], filename: str = None) -> str:
        """
        Export the report to PDF format.
        
        Args:
            report: The report content
            citations: List of citation dictionaries
            filename: Optional filename (without extension). If not provided, will use timestamp.
            
        Returns:
            Path to the exported PDF file
        """
        try:
            # Generate filename if not provided
            if filename is None:
                filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Create PDF document in exports directory
            output_path = self.exports_dir / f"{filename}.pdf"
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Create content
            content = []
            
            # Extract title and content
            title = "Research Report"
            if report.startswith("#"):
                title = report.split('\n')[0].replace('#', '').strip()
                report = '\n'.join(report.split('\n')[1:])
            
            # Add title
            content.append(Paragraph(title, self.styles["Title"]))
            content.append(Spacer(1, 24))
            
            # Process and add report content
            paragraphs = self._process_markdown(report)
            for para in paragraphs:
                if para.startswith('##'):
                    # This is a heading
                    heading = para.replace('##', '').strip()
                    content.append(Paragraph(heading, self.styles["CustomHeading1"]))
                elif para.startswith('###'):
                    # This is a subheading
                    subheading = para.replace('###', '').strip()
                    content.append(Paragraph(subheading, self.styles["CustomHeading2"]))
                else:
                    # This is normal text
                    content.append(Paragraph(para, self.styles["CustomNormal"]))
                content.append(Spacer(1, 6))
            
            # Add citations section
            if citations:
                content.append(PageBreak())
                content.append(Paragraph("References", self.styles["CustomHeading1"]))
                content.append(Spacer(1, 12))
                
                for citation in citations:
                    citation_text = f"{citation.get('title', 'Untitled')}"
                    if citation.get("url"):
                        citation_text += f" - {citation['url']}"
                    if citation.get("authors"):
                        citation_text += f" by {citation['authors']}"
                    content.append(Paragraph(citation_text, self.styles["CustomCitation"]))
                    content.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(content)
            
            logger.info(f"Exported PDF to {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error exporting PDF: {str(e)}")
            raise
    
    def export_docx(self, report: Dict[str, Any], filename: str) -> str:
        """
        Export the report to DOCX format.
        
        Args:
            report: The report dictionary
            filename: Base filename for the export
            
        Returns:
            Path to the exported DOCX file
        """
        try:
            # Create a new document
            doc = DocxDocument()
            
            # Add title
            doc.add_heading(report.get('title', 'Untitled Report'), 0)
            
            # Add sections
            for section in report.get('sections', []):
                # Add section title
                doc.add_heading(section.get('section_title', 'Untitled Section'), 1)
                
                # Add section content
                if 'content' in section and section['content']:
                    doc.add_paragraph(section['content'])
                
                # Add subsections
                if 'subsections' in section:
                    for subsection in section['subsections']:
                        # Add subsection title
                        doc.add_heading(subsection.get('subsection_title', 'Untitled Subsection'), 2)
                        
                        # Add subsection content
                        if 'content' in subsection and subsection['content']:
                            doc.add_paragraph(subsection['content'])
            
            # Add glossary if it exists
            if report.get('glossary'):
                doc.add_heading('Glossary', 1)
                for entry in report['glossary']:
                    doc.add_paragraph(f"{entry['term']}: {entry['definition']}")
            
            # Add references
            if report.get('references'):
                doc.add_heading('References', 1)
                for ref in report['references']:
                    doc.add_paragraph(f"[{ref['number']}] {ref['text']}")
            
            # Save the document
            output_path = os.path.join(EXPORTS_DIR, f"{filename}.docx")
            doc.save(output_path)
            
            logger.info(f"Exported DOCX to {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting DOCX: {str(e)}")
            raise
    
    def export_html(self, report: Dict[str, Any], filename: str) -> str:
        """
        Export the report as an HTML file.
        
        Args:
            report: Dictionary containing the report content
            filename: Name of the output file
            
        Returns:
            Path to the exported file
        """
        try:
            # Convert markdown to HTML
            html_content = markdown.markdown(report["content"])
            
            # Add citations
            if report.get("citations"):
                html_content += "\n\n<h2>References</h2>\n\n"
                for citation in report["citations"]:
                    html_content += f"<p>{citation}</p>\n\n"
            
            # Save the HTML file
            output_path = self.exports_dir / f"{filename}.html"
            with open(output_path, "w") as f:
                f.write(html_content)
            
            logger.info(f"Exported HTML to {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error exporting HTML: {str(e)}")
            raise

    def _format_report_content(self, report: Dict[str, Any]) -> str:
        """Format report content for export."""
        content = []
        
        # Add title
        content.append(f"# {report['title']}\n")
        
        # Add date
        if 'date' in report:
            content.append(f"*Generated on: {report['date']}*\n")
        
        # Process sections
        for section in report['sections']:
            # Add section title
            content.append(f"\n## {section['section_title']}\n")
            
            # Add section content
            if section.get('content'):
                content.append(f"{section['content']}\n")
            
            # Process subsections
            if section.get('subsections'):
                for subsection in section['subsections']:
                    # Add subsection title
                    content.append(f"\n### {subsection['subsection_title']}\n")
                    
                    # Add subsection content
                    if subsection.get('content'):
                        content.append(f"{subsection['content']}\n")
        
        # Add citations if available
        if any(section.get('citations') for section in report['sections']):
            content.append("\n## References\n")
            all_citations = []
            for section in report['sections']:
                all_citations.extend(section.get('citations', []))
            
            # Remove duplicates while maintaining order
            unique_citations = []
            for citation in all_citations:
                if citation not in unique_citations:
                    unique_citations.append(citation)
            
            for i, citation in enumerate(unique_citations, 1):
                content.append(f"{i}. {citation}\n")
        
        return "\n".join(content)

# Example usage
if __name__ == "__main__":
    exporter = ReportExporter()
    
    # Example report
    report = {
        "title": "Test Report",
        "content": "# Test Report\n\nThis is a test report.",
        "citations": [
            "Author, A. (2023). Test Paper. Journal of Testing, 1(1), 1-10."
        ]
    }
    
    # Export in different formats
    pdf_path = exporter.export_pdf(report["content"], report["citations"], "test_report")
    docx_path = exporter.export_docx(report, "test_report")
    html_path = exporter.export_html(report, "test_report")
    
    print(f"Exported to:\nPDF: {pdf_path}\nDOCX: {docx_path}\nHTML: {html_path}") 